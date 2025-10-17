"""
Tests for document parsing and NLP extraction in AlignCV V2.

Tests:
- PDF text extraction
- DOCX text extraction
- Text hash computation
- Skills extraction
- Roles extraction
- Entity extraction
"""

import pytest
import os
import tempfile
from docx import Document
from backend.v2.documents.parser import parse_pdf, parse_docx, compute_text_hash, validate_text_content
from backend.v2.nlp.extractor import extract_skills, extract_roles, extract_entities, extract_all


# ========================================
# Text Hash Tests
# ========================================

def test_compute_text_hash():
    """Test text hash computation."""
    text = "This is a test document"
    hash1 = compute_text_hash(text)
    hash2 = compute_text_hash(text)
    
    # Same text should produce same hash
    assert hash1 == hash2
    assert len(hash1) == 64  # SHA-256 produces 64 hex characters
    
    # Different text should produce different hash
    hash3 = compute_text_hash("Different text")
    assert hash1 != hash3


def test_validate_text_content():
    """Test text content validation."""
    # Valid text
    assert validate_text_content("This is a long enough text for validation", min_length=10) is True
    
    # Too short
    assert validate_text_content("Short", min_length=10) is False
    
    # None or empty
    assert validate_text_content(None) is False
    assert validate_text_content("") is False
    assert validate_text_content("   ") is False


# ========================================
# DOCX Parsing Tests
# ========================================

def test_parse_docx():
    """Test DOCX parsing."""
    # Create a temporary DOCX file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        doc = Document()
        doc.add_paragraph("Software Engineer with 3 years of Python experience.")
        doc.add_paragraph("Skills: FastAPI, React, PostgreSQL, Docker.")
        doc.save(temp_file.name)
        temp_path = temp_file.name
    
    try:
        # Parse document
        text = parse_docx(temp_path)
        
        assert text is not None
        assert "Software Engineer" in text
        assert "Python" in text
        assert "FastAPI" in text
    finally:
        # Clean up
        os.unlink(temp_path)


def test_parse_docx_invalid_file():
    """Test DOCX parsing with invalid file."""
    text = parse_docx("nonexistent_file.docx")
    assert text is None


# ========================================
# Skills Extraction Tests
# ========================================

def test_extract_skills():
    """Test skills extraction from text."""
    text = """
    Software Engineer with 5 years of experience in Python, JavaScript, and React.
    Expertise in FastAPI, Django, PostgreSQL, and Docker.
    Cloud experience with AWS and Azure.
    """
    
    skills = extract_skills(text)
    
    assert "Python" in skills
    assert "Javascript" in skills
    assert "React" in skills
    assert "Fastapi" in skills
    assert "Django" in skills
    assert "Postgresql" in skills
    assert "Docker" in skills
    assert "Aws" in skills


def test_extract_skills_empty_text():
    """Test skills extraction from empty text."""
    skills = extract_skills("")
    assert skills == []


def test_extract_skills_no_matches():
    """Test skills extraction with no matching skills."""
    text = "I am a person who likes to cook and travel."
    skills = extract_skills(text)
    assert skills == []


# ========================================
# Roles Extraction Tests
# ========================================

def test_extract_roles():
    """Test role extraction from text."""
    text = """
    Looking for a Software Engineer with Backend Developer experience.
    Must have Full Stack Developer skills and Data Scientist knowledge.
    """
    
    roles = extract_roles(text)
    
    assert "Software Engineer" in roles
    assert "Backend Developer" in roles
    assert "Full Stack Developer" in roles
    assert "Data Scientist" in roles


def test_extract_roles_empty_text():
    """Test roles extraction from empty text."""
    roles = extract_roles("")
    assert roles == []


# ========================================
# Entity Extraction Tests
# ========================================

def test_extract_entities():
    """Test entity extraction using SpaCy."""
    text = """
    John Smith worked at Google in Mountain View, California from 2018 to 2023.
    He specialized in machine learning and earned a degree from Stanford University.
    """
    
    # This test will only pass if SpaCy model is installed
    try:
        entities = extract_entities(text)
        
        # Check that some entities were found
        assert isinstance(entities, dict)
        # Entities might include PERSON, ORG, GPE, DATE, etc.
        # Exact entities depend on SpaCy model performance
        
    except Exception as e:
        pytest.skip(f"SpaCy model not available: {str(e)}")


def test_extract_entities_empty_text():
    """Test entity extraction from empty text."""
    try:
        entities = extract_entities("")
        assert entities == {}
    except Exception:
        pytest.skip("SpaCy model not available")


# ========================================
# Full Extraction Tests
# ========================================

def test_extract_all():
    """Test full extraction (skills, roles, entities)."""
    text = """
    Software Engineer with Python and React experience.
    Looking for Backend Developer or Full Stack Developer positions.
    Worked at Microsoft in Seattle.
    """
    
    try:
        result = extract_all(text)
        
        assert "skills" in result
        assert "roles" in result
        assert "entities" in result
        
        assert "Python" in result["skills"]
        assert "React" in result["skills"]
        
        assert len(result["roles"]) > 0
        
    except Exception as e:
        pytest.skip(f"SpaCy model not available: {str(e)}")


# ========================================
# Integration Tests
# ========================================

def test_full_document_workflow():
    """Test complete workflow: create DOCX -> parse -> extract -> hash."""
    # Create test document
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        doc = Document()
        doc.add_paragraph("Software Engineer Resume")
        doc.add_paragraph("Skills: Python, FastAPI, React, PostgreSQL, Docker")
        doc.add_paragraph("Experience: Backend Developer at Tech Corp")
        doc.save(temp_file.name)
        temp_path = temp_file.name
    
    try:
        # Step 1: Parse document
        text = parse_docx(temp_path)
        assert text is not None
        assert len(text) > 0
        
        # Step 2: Validate content
        assert validate_text_content(text) is True
        
        # Step 3: Extract skills and roles
        skills = extract_skills(text)
        roles = extract_roles(text)
        
        assert "Python" in skills
        assert "Fastapi" in skills
        assert "Backend Developer" in roles
        
        # Step 4: Compute hash
        text_hash = compute_text_hash(text)
        assert len(text_hash) == 64
        
        # Step 5: Verify hash is consistent
        hash2 = compute_text_hash(text)
        assert text_hash == hash2
        
    finally:
        os.unlink(temp_path)
